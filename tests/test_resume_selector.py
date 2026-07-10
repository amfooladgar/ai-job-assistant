import pytest
import os
import hashlib
from pathlib import Path
from src.storage.db import JobDatabase
from src.sources.base import JobPosting
from src.agents.resume_selector_agent import ResumeSelectorAgent

def test_resume_profiling_and_caching(tmp_path):
    # Setup temp resumes directory
    resumes_dir = tmp_path / "resumes"
    resumes_dir.mkdir()
    
    resume_text = (
        "Alex Reed\n"
        "Machine Learning Engineer\n"
        "Expertise in PyTorch, NLP, JAX, Python"
    )
    resume_file = resumes_dir / "alex_ml.txt"
    with open(resume_file, "w", encoding="utf-8") as f:
        f.write(resume_text)
        
    db_path = tmp_path / "test_selector.db"
    db = JobDatabase(db_path)
    
    agent = ResumeSelectorAgent(
        resumes_dir=resumes_dir,
        db=db,
        enable_llm_reasoning=False
    )
    
    content_hash = hashlib.sha256(resume_text.encode("utf-8")).hexdigest()
    
    # 1. Verify no cache initially
    cached = db.get_cached_resume("alex_ml.txt", content_hash)
    assert cached is None
    
    # 2. Run selection to trigger profiling and caching
    job = JobPosting(
        title="ML Engineer",
        company="AI Inc",
        location="Remote",
        description="We need someone who writes PyTorch and understands NLP.",
        url="https://example.com/ml-job",
        source="test"
    )
    
    selected_file, content, reasoning = agent.select_best_resume(job)
    assert selected_file == "alex_ml.txt"
    assert content == resume_text
    
    # 3. Verify it is now cached in the database
    cached = db.get_cached_resume("alex_ml.txt", content_hash)
    assert cached is not None
    assert "PyTorch" in cached["skills"]["frameworks_libraries"]
    
    # 4. Modify resume and verify hash mismatch results in cache invalidation
    new_resume_text = resume_text + "\nAlso knows HTML and CSS."
    with open(resume_file, "w", encoding="utf-8") as f:
        f.write(new_resume_text)
        
    # The old hash should not match the new file contents
    new_hash = hashlib.sha256(new_resume_text.encode("utf-8")).hexdigest()
    
    # Cache for new hash shouldn't exist yet
    assert db.get_cached_resume("alex_ml.txt", new_hash) is None
    
    # Running selector again triggers re-profiling for new content
    selected_file2, content2, reasoning2 = agent.select_best_resume(job)
    assert selected_file2 == "alex_ml.txt"
    assert content2 == new_resume_text
    
    # Verify new hash is cached
    assert db.get_cached_resume("alex_ml.txt", new_hash) is not None

def test_resume_selection_matching(tmp_path):
    resumes_dir = tmp_path / "resumes"
    resumes_dir.mkdir()
    
    # Create ML resume
    ml_text = "Alex Reed\nML Engineer\nPyTorch, Python, NLP, LLMs"
    with open(resumes_dir / "ml.txt", "w", encoding="utf-8") as f:
        f.write(ml_text)
        
    # Create Frontend resume
    fe_text = "John Smith\nFrontend Developer\nReact, TypeScript, CSS, Javascript"
    with open(resumes_dir / "frontend.txt", "w", encoding="utf-8") as f:
        f.write(fe_text)
        
    db = JobDatabase(tmp_path / "test_selector_match.db")
    agent = ResumeSelectorAgent(resumes_dir=resumes_dir, db=db, enable_llm_reasoning=False)
    
    # Test job 1: ML Engineer
    job_ml = JobPosting(
        title="Machine Learning Engineer",
        company="AI Labs",
        location="Remote",
        description="Join us to train large language models in Python and PyTorch.",
        url="https://example.com/ml",
        source="test"
    )
    selected_file, _, reasoning = agent.select_best_resume(job_ml)
    assert selected_file == "ml.txt"
    
    # Test job 2: React Developer
    job_fe = JobPosting(
        title="Web Developer",
        company="Web Corp",
        location="Remote",
        description="Building user interfaces using React, typescript and Tailwind CSS.",
        url="https://example.com/fe",
        source="test"
    )
    selected_file2, _, reasoning2 = agent.select_best_resume(job_fe)
    assert selected_file2 == "frontend.txt"

def test_resume_formats_pdf_docx(tmp_path):
    import docx
    from unittest.mock import MagicMock, patch
    
    resumes_dir = tmp_path / "resumes"
    resumes_dir.mkdir()
    
    # 1. Create a DOCX resume
    doc = docx.Document()
    doc.add_paragraph("Alex Reed")
    doc.add_paragraph("Frontend Developer")
    doc.add_paragraph("Skills: React, TypeScript, CSS, Javascript")
    docx_path = resumes_dir / "resume.docx"
    doc.save(docx_path)
    
    # 2. Mock a PDF reader response since creating binary PDFs from scratch is complex
    mock_pdf_page = MagicMock()
    mock_pdf_page.extract_text = MagicMock(return_value="Alex Reed\nML Engineer\nPyTorch, Python, NLP, LLMs")
    
    mock_pdf_reader = MagicMock()
    mock_pdf_reader.pages = [mock_pdf_page]
    
    db = JobDatabase(tmp_path / "test_formats.db")
    agent = ResumeSelectorAgent(resumes_dir=resumes_dir, db=db, enable_llm_reasoning=False)
    
    with patch("pypdf.PdfReader", return_value=mock_pdf_reader):
        # We simulate that a pdf file exists
        pdf_path = resumes_dir / "resume.pdf"
        with open(pdf_path, "w") as f:
            f.write("mock binary pdf content")
            
        # Test loading DOCX content
        docx_content = agent._read_file_content(docx_path)
        assert "React" in docx_content
        assert "TypeScript" in docx_content
        
        # Test loading PDF content
        pdf_content = agent._read_file_content(pdf_path)
        assert "PyTorch" in pdf_content
        assert "LLMs" in pdf_content

